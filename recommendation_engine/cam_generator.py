from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import json
import re
from config import OUTPUT_FOLDER


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_table(doc: Document, data: dict, title: str = ""):
    if title:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Parameter"
    table.rows[0].cells[1].text = "Value"
    for key, value in data.items():
        row = table.add_row()
        row.cells[0].text = str(key).replace("_", " ").title()
        row.cells[1].text = str(value)
    doc.add_paragraph()


def format_inr(value) -> str:
    try:
        num = float(value)
        if num >= 10000000:
            return f"₹{num/10000000:.2f} Cr"
        elif num >= 100000:
            return f"₹{num/100000:.2f} L"
        else:
            return f"₹{num:,.0f}"
    except:
        return str(value)


def format_financial_data(financial_data: dict) -> dict:
    formatted = {}
    rupee_fields = ["loan_amount_requested", "net_worth", "revolving_balance",
                    "total_current_balance", "total_revolving_credit_limit"]
    pct_fields = ["revenue_growth", "revolving_utilities", "interest_rate"]
    ratio_fields = ["debt_service_coverage", "current_ratio", "debt_to_equity",
                    "collateral_coverage"]

    for key, value in financial_data.items():
        label = key.replace("_", " ").title()
        if key in rupee_fields:
            formatted[label] = format_inr(value)
        elif key in pct_fields:
            try:
                pct = float(value) * 100 if float(value) < 1 else float(value)
                formatted[label] = f"{pct:.1f}%"
            except:
                formatted[label] = str(value)
        elif key in ratio_fields:
            try:
                formatted[label] = f"{float(value):.2f}x"
            except:
                formatted[label] = str(value)
        else:
            formatted[label] = str(value)
    return formatted


def parse_research(analysis) -> list:
    import ast
    clean = re.sub(r"```json|```", "", str(analysis)).strip()
    try:
        parsed = json.loads(clean)
        items = []
        if isinstance(parsed, dict):
            for key, value in parsed.items():
                if isinstance(value, dict):
                    for subkey, subval in value.items():
                        # Clean list formatting
                        if isinstance(subval, list):
                            subval = " | ".join([str(v) for v in subval])
                        elif isinstance(subval, str) and subval.startswith("["):
                            try:
                                parsed_list = ast.literal_eval(subval)
                                subval = " | ".join([str(v) for v in parsed_list])
                            except:
                                subval = subval.strip("[]'\"")
                        items.append((
                            f"{key.replace('_', ' ').title()} — {subkey.replace('_', ' ').title()}",
                            str(subval)[:800]
                        ))
                else:
                    # Clean list formatting
                    if isinstance(value, list):
                        value = " | ".join([str(v) for v in value])
                    elif isinstance(value, str) and value.startswith("["):
                        try:
                            parsed_list = ast.literal_eval(value)
                            value = " | ".join([str(v) for v in parsed_list])
                        except:
                            value = value.strip("[]'\"")
                    items.append((key.replace("_", " ").title(), str(value)[:800]))
        return items
    except:
        return [("Research Findings", clean[:1000])]


def generate_cam(
    company_name: str,
    financial_data: dict,
    gst_analysis: dict,
    research_analysis: dict,
    five_cs: dict,
    decision: dict,
    explanation: str,
    qualitative_notes: list = [],
    swot=None
) -> str:
    doc = Document()

    # Title
    title = doc.add_heading("CREDIT APPRAISAL MEMO (CAM)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Borrower: {company_name}  |  Date: {datetime.now().strftime('%d %B %Y')}  |  Prepared by: Intelli-Credit Engine"
    )
    doc.add_paragraph()

    # Risk Rating Banner
    risk_map = {
        "APPROVE": "OVERALL RATING: MEDIUM RISK ⚠️ — APPROVED WITH CONDITIONS",
        "CONDITIONAL APPROVE": "OVERALL RATING: MEDIUM RISK ⚠️ — CONDITIONAL APPROVAL",
        "REJECT": "OVERALL RATING: HIGH RISK ❌ — REJECTED"
    }
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run(risk_map.get(decision["decision"], "OVERALL RATING: MEDIUM RISK ⚠️"))
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    # Company Overview
    add_heading(doc, "Company Overview")
    overview_data = {
        "Company Name": company_name,
        "Loan Requested": format_inr(financial_data.get("loan_amount_requested", 0)),
        "Assessment Date": datetime.now().strftime("%d %B %Y"),
        "Sector Outlook": financial_data.get("sector_outlook", "N/A").title(),
        "Collateral Type": financial_data.get("collateral_type", "N/A").title(),
        "Net Worth": format_inr(financial_data.get("net_worth", 0)),
    }
    add_table(doc, overview_data)

    # Section 1
    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(explanation)

    # Section 2
    add_heading(doc, "2. Credit Decision")
    add_table(doc, {
        "Decision": decision["decision"],
        "Credit Limit": format_inr(decision["credit_limit"]),
        "Risk Premium": f"{decision['risk_premium']}%" if decision["risk_premium"] else "N/A",
        "Credit Score": f"{decision['score']}/100",
        "ML Model Score": f"{decision.get('ml_score', 'N/A')}/100",
    })

    # Section 3
    add_heading(doc, "3. Five Cs of Credit Analysis")
    descriptions = {
        "character": "Promoter reputation, litigation history, external research findings",
        "capacity": "Debt service coverage ratio, cash flows, revenue growth trajectory",
        "capital": "Net worth position, debt-to-equity ratio, financial leverage",
        "collateral": "Asset coverage ratio, collateral type and quality of security",
        "conditions": "Sector outlook, GST compliance health, macroeconomic conditions"
    }
    for c, score in five_cs.items():
        rating = "Strong" if score >= 70 else "Adequate" if score >= 50 else "Weak"
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{c.title()} ({score}/100 — {rating}): ").bold = True
        p.add_run(descriptions.get(c, ""))
    doc.add_paragraph()

    # Section 4
    add_heading(doc, "4. Financial Analysis")
    add_table(doc, format_financial_data(financial_data))

    # Section 5
    add_heading(doc, "5. GST & Bank Statement Analysis")
    add_table(doc, {
        "GST Revenue": format_inr(gst_analysis.get("gst_revenue", 0)),
        "Bank Credits": format_inr(gst_analysis.get("bank_credits", 0)),
        "Discrepancy": f"{gst_analysis.get('discrepancy_pct', 0)}%",
        "ITC Compliance": "GSTR-2B Analysis — " + ("High Risk" if gst_analysis.get("risk_level") == "high" else "Compliant"),
        "Risk Level": gst_analysis.get("risk_level", "N/A").upper(),
        "Fraud Flags": ", ".join(gst_analysis.get("flags", [])) or "None detected"
    })

    # Section 6
    add_heading(doc, "6. External Research & Due Diligence")
    research_items = parse_research(research_analysis.get("analysis", ""))
    for label, content in research_items:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(content)
    if research_analysis.get("sources"):
        doc.add_paragraph()
        src_para = doc.add_paragraph()
        src_para.add_run("References & Data Sources:").bold = True
        src_para.add_run(", ".join(research_analysis["sources"][:5]))

    # Section 7
    add_heading(doc, "7. Qualitative Observations (Credit Officer)")
    if qualitative_notes:
        for note in qualitative_notes:
            doc.add_paragraph(note, style="List Bullet")
    else:
        doc.add_paragraph("No qualitative observations recorded.")

    # Section 8
    add_heading(doc, "8. Disclaimer")
    doc.add_paragraph(
        "This CAM has been auto-generated by the Intelli-Credit Engine using AI and ML models. "
        "It is intended to assist credit officers in decision-making and does not replace human judgment. "
        "Final lending decisions must be approved by an authorized credit committee as per RBI guidelines."
    )

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    filename = f"{OUTPUT_FOLDER}/CAM_{company_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename